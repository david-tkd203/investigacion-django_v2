import os
from pathlib import Path
from django.conf import settings
from docx import Document


def export_docx_wrapper(accidente, informe, resumen_texto: str = "") -> str:
    """
    Genera un DOCX y devuelve la ruta absoluta del archivo en PROTECTED_MEDIA.
    Si tienes tu propio wrapper (el de Streamlit), llama aquí a ese wrapper
    y retorna el path final.

    Returns: str -> ruta absoluta del docx generado
    """
    # Carpeta destino protegida
    base = Path(settings.PROTECTED_MEDIA_ROOT) / "informes" / accidente.codigo_accidente
    base.mkdir(parents=True, exist_ok=True)
    filename = f"{informe.codigo}_v{informe.version}.docx"
    path = base / filename

    #  Minimalista: genera un docx simple si python-docx está instalado
    try:
        doc = Document()
        doc.add_heading('Informe de Investigación de Accidente', level=1)
        doc.add_paragraph(f"Informe N°: {informe.codigo}")
        doc.add_paragraph(f"Accidente: {accidente.codigo_accidente}")
        doc.add_paragraph(f"Investigador: {informe.investigador}")
        doc.add_paragraph(f"Fecha Informe: {informe.fecha_informe.strftime('%Y-%m-%d')}")
        if resumen_texto:
            doc.add_heading('Resumen', level=2)
            doc.add_paragraph(resumen_texto)
        doc.save(path.as_posix())
    except Exception:
        # Si no está python-docx, crea un .txt como placeholder
        path = path.with_suffix(".txt")
        path.write_text(
            f"Informe N°: {informe.codigo}\n"
            f"Accidente: {accidente.codigo_accidente}\n"
            f"Investigador: {informe.investigador}\n"
            f"Fecha Informe: {informe.fecha_informe:%Y-%m-%d}\n"
            f"{'-'*30}\n{resumen_texto or '(sin resumen)'}",
            encoding="utf-8"
        )

    return path.as_posix()
