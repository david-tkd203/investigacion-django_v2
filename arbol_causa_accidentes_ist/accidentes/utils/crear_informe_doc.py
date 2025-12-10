# -*- coding: utf-8 -*-
from __future__ import annotations
import logging
from pathlib import Path
from datetime import date, datetime
from typing import Any, Dict, Iterable, Tuple, Optional, List

from django.conf import settings
from django.apps import apps

log = logging.getLogger(__name__)


class InformeDocxBuilder:
    """
    Construye el archivo de informe (DOCX si python-docx está disponible;
    TXT como fallback). Expone un único método público build().
    """

    # ---------- Helpers estáticos seguros ----------
    @staticmethod
    def _fmt_date(d: Any) -> str:
        if isinstance(d, (date, datetime)):
            d2 = d if isinstance(d, date) else d.date()
            return d2.strftime("%d/%m/%Y")
        try:
            return datetime.fromisoformat(str(d)).strftime("%d/%m/%Y")
        except Exception:
            return str(d) if d else ""

    @staticmethod
    def _getattr_safe(obj: Any, path: str, default: str = "") -> str:
        cur = obj
        try:
            for part in path.split("."):
                if cur is None:
                    return default
                cur = getattr(cur, part)
            return "" if cur is None else str(cur)
        except Exception:
            return default

    @staticmethod
    def _get_model_if_exists(app_label: str, model_name: str):
        try:
            return apps.get_model(app_label, model_name)
        except Exception:
            return None

    @staticmethod
    def _first_non_empty(*values: Any) -> str:
        for v in values:
            if v is not None and str(v).strip():
                return str(v)
        return ""

    @staticmethod
    def _fmt_aym(anos: object, meses: object) -> str:
        """
        Formatea años/meses a texto: 'X años Y meses'.
        Oculta las partes en 0 y usa '0 meses' si ambos son 0/None.
        """
        try:
            y = int(anos or 0)
        except Exception:
            y = 0
        try:
            m = int(meses or 0)
        except Exception:
            m = 0

        # seguridad en rango de meses
        m = max(0, min(11, m))

        partes = []
        if y:
            partes.append(f"{y} año" if y == 1 else f"{y} años")
        if m:
            partes.append(f"{m} mes" if m == 1 else f"{m} meses")

        if not partes:
            return "0 meses"
        return " ".join(partes)

    # ---------- Lectura de datos ----------
    def _leer_empresa_y_centro(self, accidente) -> Dict[str, str]:
        emp = None
        try:
            emp = getattr(accidente.trabajador, "empresa", None)
        except Exception:
            pass
        if emp is None:
            try:
                emp = getattr(accidente.centro, "empresa", None)
            except Exception:
                pass

        centro = getattr(accidente, "centro", None)

        def g(obj, attr, default=""):
            return getattr(obj, attr, default) if obj else default

        return {
            "razon_social":       self._first_non_empty(g(emp, "empresa_sel")),
            "rut_empresa":        self._first_non_empty(g(emp, "rut_empresa")),
            "actividad":          self._first_non_empty(g(emp, "actividad")),
            "direccion_empresa":  self._first_non_empty(g(emp, "direccion_empresa")),
            "telefono_empresa":   self._first_non_empty(g(emp, "telefono")),
            "representante_legal":self._first_non_empty(g(emp, "representante_legal")),
            "region":             self._first_non_empty(g(emp, "region"), g(centro, "region")),
            "comuna":             self._first_non_empty(g(emp, "comuna"), g(centro, "comuna")),
            "centro_trabajo":     self._first_non_empty(g(centro, "nombre_local")),
            "direccion_centro":   self._first_non_empty(g(centro, "direccion_centro")),
        }

    def _leer_trabajador(self, accidente) -> Dict[str, str]:
        t = getattr(accidente, "trabajador", None)

        def g(attr, default=None):
            return getattr(t, attr, default) if t else default

        # Nuevos campos del modelo (años/meses)
        ae_y = g("antiguedad_empresa_anios", 0)
        ae_m = g("antiguedad_empresa_meses", 0)
        ac_y = g("antiguedad_cargo_anios", 0)
        ac_m = g("antiguedad_cargo_meses", 0)

        out = {
            "nombre":             g("nombre_trabajador", "") or "",
            "rut":                g("rut_trabajador", "") or "",
            "fecha_nacimiento":   self._fmt_date(g("fecha_nacimiento", "")) or "",
            "edad":               "",
            "nacionalidad":       g("nacionalidad", "") or "",
            "estado_civil":       g("estado_civil", "") or "",
            "tipo_contrato":      g("contrato", "") or "",
            # Mantener las claves históricas pero con el nuevo formato A/M
            "antiguedad_empresa": self._fmt_aym(ae_y, ae_m),
            "cargo":              g("cargo_trabajador", "") or "",
            "antiguedad_cargo":   self._fmt_aym(ac_y, ac_m),
            "domicilio":          g("domicilio", "") or "",
        }

        # Calcular edad si hay fecha de nacimiento
        if not out["edad"] and out["fecha_nacimiento"]:
            try:
                b = datetime.strptime(out["fecha_nacimiento"], "%d/%m/%Y").date()
                t_today = date.today()
                out["edad"] = str(t_today.year - b.year - ((t_today.month, t_today.day) < (b.month, b.day)))
            except Exception:
                pass
        return out

    def _leer_datos_accidente(self, accidente) -> Dict[str, str]:
        def g(attr, default=""):
            return getattr(accidente, attr, default)
        return {
            "fecha":            self._fmt_date(g("fecha_accidente")),
            "hora":             self._first_non_empty(g("hora_accidente")),
            "lugar":            self._first_non_empty(g("lugar_accidente")),
            "tipo":             self._first_non_empty(g("tipo_accidente")),
            "naturaleza":       self._first_non_empty(g("naturaleza_lesion")),
            "parte_afectada":   self._first_non_empty(g("parte_afectada")),
            "tarea":            self._first_non_empty(g("tarea")),
            "operacion":        self._first_non_empty(g("operacion")),
            "danos_personas":   self._first_non_empty(g("danos_personas")),
            "danos_propiedad":  self._first_non_empty(g("danos_propiedad")),
            "perdidas_proceso": self._first_non_empty(g("perdidas_proceso")),
        }

    def _leer_hechos(self, accidente) -> Iterable[str]:
        Model = self._get_model_if_exists("accidentes", "Hechos")
        if not Model:
            return []
        try:
            qs = Model.objects.filter(accidente_id=accidente.accidente_id).order_by("secuencia", "hecho_id")
            textos = []
            for h in qs:
                txt = self._first_non_empty(getattr(h, "descripcion", None), str(h))
                if txt.strip():
                    textos.append(txt.strip())
            return textos
        except Exception:
            return []

    def _leer_prescripciones(self, accidente):
        Model = self._get_model_if_exists("accidentes", "Prescripciones")
        if not Model:
            return []
        try:
            qs = Model.objects.filter(accidente_id=accidente.accidente_id).order_by("prescripcion_id")
            out = []
            for m in qs:
                out.append({
                    "tipo":        self._first_non_empty(getattr(m, "tipo", None)),
                    "prioridad":   self._first_non_empty(getattr(m, "prioridad", None)),
                    "plazo":       self._fmt_date(getattr(m, "plazo", None)),
                    "responsable": self._first_non_empty(getattr(m, "responsable", None)),
                    "descripcion": self._first_non_empty(getattr(m, "descripcion", None), str(m)),
                })
            return out
        except Exception:
            return []

    def _leer_declaraciones(self, accidente, max_items: int = 6) -> List[Dict[str, str]]:
        Model = self._get_model_if_exists("accidentes", "Declaraciones")
        if not Model:
            return []
        try:
            qs = Model.objects.filter(accidente_id=accidente.accidente_id).order_by("declaracion_id")[:max_items]
            out = []
            for d in qs:
                out.append({
                    "tipo":   getattr(d, "tipo_decl", "") or "",
                    "nombre": getattr(d, "nombre", "") or "",
                    "cargo":  getattr(d, "cargo", "") or "",
                    "rut":    getattr(d, "rut", "") or "",
                    "texto":  (getattr(d, "texto", "") or "")[:800],
                })
            return out
        except Exception:
            return []

    def _leer_documentos(self, accidente, max_items: int = 10) -> List[Dict[str, str]]:
        Model = self._get_model_if_exists("accidentes", "Documentos")
        if not Model:
            return []
        try:
            qs = Model.objects.filter(accidente_id=accidente.accidente_id).order_by("-subido_el")[:max_items]
            out = []
            for d in qs:
                out.append({
                    "nombre":  getattr(d, "documento", "") or getattr(d, "nombre_archivo", "") or "",
                    "objetivo":getattr(d, "objetivo", "") or "",
                    "mime":    getattr(d, "mime_type", "") or "",
                    "fecha":   self._fmt_date(getattr(d, "subido_el", "")),
                    "url":     getattr(d, "url_descarga", "#"),
                })
            return out
        except Exception:
            return []

    # ---------- Relato (manejo con is_current) ----------
    def _get_current_relato_record(self, accidente):
        ModelRel = (
            self._get_model_if_exists("accidentes", "Relato")
            or self._get_model_if_exists("accidentes", "AccidentesRelato")
            or self._get_model_if_exists("accidentes", "Relatos")
        )
        if not ModelRel:
            return None

        qs = ModelRel.objects.filter(accidente_id=accidente.accidente_id)
        field_names = {f.name for f in ModelRel._meta.fields}

        if "is_current" in field_names:
            qs = qs.filter(is_current=True)

        order = []
        if "version" in field_names:
            order.append("-version")
        order.append("-pk")
        return qs.order_by(*order).first()

    def _obtener_relato_final(self, accidente) -> str:
        try:
            rel = self._get_current_relato_record(accidente)
            if not rel:
                return ""
            for f in ("relato_final", "relato_inicial", "texto", "contenido"):
                if hasattr(rel, f) and getattr(rel, f):
                    return str(getattr(rel, f)).strip()
        except Exception:
            pass
        return ""

    def _obtener_relato_para_resumen(self, accidente) -> str:
        try:
            rel = self._get_current_relato_record(accidente)
            if rel:
                piezas = []
                for f in ("relato_final", "relato_inicial",
                          "respuesta_1", "respuesta_2", "respuesta_3",
                          "texto", "contenido"):
                    if hasattr(rel, f) and getattr(rel, f):
                        piezas.append(str(getattr(rel, f)))
                if piezas:
                    return "\n\n".join(piezas)[:6000]
        except Exception:
            pass

        datos = self._leer_datos_accidente(accidente)
        piezas_fb = [
            f"Fecha: {datos.get('fecha','')}",
            f"Hora: {datos.get('hora','')}",
            f"Lugar: {datos.get('lugar','')}",
            f"Contexto: {self._first_non_empty(self._getattr_safe(accidente, 'contexto'), '')}",
            f"Circunstancias: {self._first_non_empty(self._getattr_safe(accidente, 'circunstancias'), '')}",
        ]
        return "\n".join(p for p in piezas_fb if p.strip())

    def _resumen_via_ia(self, relato_texto_or_accidente) -> str:
        """
        Si hay 'accidentes.resumen' no vacío, lo devuelve y no llama a IA.
        Si está vacío o no existe, genera el resumen con IA y lo guarda (<=1000 chars).
        El argumento puede ser:
        - int / str numérica: accidente_id
        - instancia con .accidente_id
        - str no numérica: texto directo a resumir
        """
        acc_id = 0
        fuente = ""

        # --- Resolver accidente_id (si vino) ---
        try:
            if hasattr(relato_texto_or_accidente, "accidente_id"):
                acc_id = int(getattr(relato_texto_or_accidente, "accidente_id") or 0)
        except Exception:
            acc_id = 0

        if not acc_id:
            try:
                if isinstance(relato_texto_or_accidente, int):
                    acc_id = int(relato_texto_or_accidente)
                elif isinstance(relato_texto_or_accidente, str) and relato_texto_or_accidente.strip().isdigit():
                    acc_id = int(relato_texto_or_accidente.strip())
            except Exception:
                acc_id = 0

        # --- 1) Si hay resumen manual en BD, devolverlo y salir ---
        if acc_id > 0:
            try:
                AccModel = self._get_model_if_exists("accidentes", "Accidentes")
                if AccModel:
                    acc_obj = AccModel.objects.filter(accidente_id=acc_id).only("resumen").first()
                    if acc_obj:
                        existing = (acc_obj.resumen or "").strip()
                        if existing:
                            # Resumen ya escrito por alguien -> respetar y no llamar a IA
                            return existing
            except Exception as e:
                log.warning("Resumen IA: no se pudo leer accidentes.resumen (accidente_id=%s): %s", acc_id, e)

        # --- 2) Armar 'fuente' (texto a resumir) ---
        if acc_id > 0:
            # Cargar relato desde tablas conocidas
            try:
                ModelRel = (
                    self._get_model_if_exists("accidentes", "AccidentesRelato")
                    or self._get_model_if_exists("accidentes", "Relato")
                    or self._get_model_if_exists("accidentes", "Relatos")
                )
                if ModelRel:
                    qs = ModelRel.objects.filter(accidente_id=acc_id)
                    field_names = {f.name for f in ModelRel._meta.fields}
                    if "is_current" in field_names:
                        qs = qs.filter(is_current=True)

                    order = []
                    if "version" in field_names:
                        order.append("-version")
                    order.append("-pk")
                    rel = qs.order_by(*order).first()

                    if rel:
                        for f in ("relato_final", "relato_inicial", "texto", "contenido"):
                            if hasattr(rel, f) and getattr(rel, f):
                                fuente = str(getattr(rel, f)).strip()
                                break
            except Exception as e:
                log.warning("Resumen IA: error leyendo relato para accidente_id=%s: %s", acc_id, e)

        # Si no hay acc_id o no se pudo leer relato, quizá vino texto directo
        if not fuente:
            if isinstance(relato_texto_or_accidente, str) and not relato_texto_or_accidente.strip().isdigit():
                fuente = relato_texto_or_accidente.strip()
            else:
                # Sin texto de entrada: nada que resumir
                return ""

        # --- 3) Llamada a IA (solo si no había resumen manual) ---
        try:
            from accidentes.views_api.prompt_utils import call_ia_text
        except Exception as e:
            log.warning("No se pudo importar prompt_utils.call_ia_text: %s", e)
            return ""

        prev_in = (fuente[:300] + "…") if len(fuente) > 300 else fuente
        log.debug("Resumen IA -> entrada (accidente_id=%s, len=%s): %r", acc_id or "-", len(fuente), prev_in)

        try:
            out = call_ia_text(fuente, prompt_key="resumen")
            resumen = (out or "").strip()
            if len(resumen) > 1000:
                resumen = resumen[:1000]

            prev_out = (resumen[:300] + "…") if len(resumen) > 300 else resumen
            log.debug("Resumen IA <- salida (accidente_id=%s, len=%s): %r", acc_id or "-", len(resumen), prev_out)
        except Exception as e:
            log.warning("Fallo llamada IA (resumen): %s", e)
            return ""

        # --- 4) Guardar en accidentes.resumen si hay accidente_id ---
        if acc_id > 0 and resumen:
            try:
                AccModel = self._get_model_if_exists("accidentes", "Accidentes")
                if AccModel:
                    from django.db import transaction
                    with transaction.atomic():
                        updated = AccModel.objects.filter(accidente_id=acc_id).update(resumen=resumen)
                        log.debug("Resumen IA: guardado en accidentes.resumen (accidente_id=%s, updated=%s)", acc_id, updated)
                        try:
                            if hasattr(relato_texto_or_accidente, "resumen"):
                                relato_texto_or_accidente.resumen = resumen
                            # Si vino la propia instancia Accidente, también intenta reflejar:
                            if hasattr(relato_texto_or_accidente, "accidente_id") and hasattr(relato_texto_or_accidente, "resumen"):
                                relato_texto_or_accidente.resumen = resumen
                        except Exception:
                            pass
            except Exception as e:
                log.warning("No se pudo guardar accidentes.resumen (accidente_id=%s): %s", acc_id, e)

        return resumen


    def _intentar_resumen_relato(self, accidente_id: int) -> str:
        try:
            ModelRel = (
                self._get_model_if_exists("accidentes", "AccidentesRelato")
                or self._get_model_if_exists("accidentes", "Relato")
                or self._get_model_if_exists("accidentes", "Relatos")
            )
            if not ModelRel:
                return ""

            qs = ModelRel.objects.filter(accidente_id=accidente_id)
            names = {f.name for f in ModelRel._meta.fields}
            if "is_current" in names:
                qs = qs.filter(is_current=True)

            order = []
            if "version" in names:
                order.append("-version")
            order.append("-pk")
            rel = qs.order_by(*order).first()
            if not rel:
                return ""

            for f in ("relato_final", "contenido", "texto", "resumen"):
                if hasattr(rel, f) and getattr(rel, f):
                    return str(getattr(rel, f))[:4000]
        except Exception:
            return ""
        return ""

    # ---------- Render del Árbol a PNG ----------
    def _render_arbol_png(self, accidente) -> Optional[str]:
        ModelArbol = self._get_model_if_exists("accidentes", "ArbolCausas")
        if not ModelArbol:
            return None

        arbol = (
            ModelArbol.objects.filter(accidente_id=accidente.accidente_id, is_current=True)
            .order_by("-version")
            .first()
        )
        if not arbol:
            return None

        dot_src = None
        if getattr(arbol, "arbol_json_dot", None):
            dot_src = arbol.arbol_json_dot
        elif getattr(arbol, "arbol_json_5q", None):
            try:
                from accidentes.utils.causal_tree import CausalTree
                tree = CausalTree(arbol_json_5q=arbol.arbol_json_5q)
                if hasattr(tree, "to_dot"):
                    dot_src = tree.to_dot()
                elif hasattr(tree, "export_dot"):
                    dot_src = tree.export_dot()
                else:
                    dot_src = str(tree)
            except Exception as e:
                log.warning("No fue posible generar DOT desde 5Q: %s", e)
                dot_src = None

        if not dot_src:
            return None

        base_dir = Path(settings.PROTECTED_MEDIA_ROOT) / "informes" / str(accidente.codigo_accidente)
        base_dir.mkdir(parents=True, exist_ok=True)
        png_path_stem = base_dir / "arbol_causas"
        try:
            from graphviz import Source
            s = Source(dot_src)
            out = s.render(filename=png_path_stem.as_posix(), format="png", cleanup=True)
            return out
        except Exception as e:
            log.warning("Fallo renderizando DOT->PNG con Graphviz: %s", e)
            try:
                (png_path_stem.with_suffix(".dot")).write_text(dot_src, encoding="utf-8")
            except Exception:
                pass
            return None

    # ---------- Helpers de formateo DOCX ----------
    def _add_kv_table(self, doc, rows: List[Tuple[str, str]]):
        from docx.shared import Inches  # noqa: F401
        table = doc.add_table(rows=len(rows), cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for i, (k, v) in enumerate(rows):
            row = table.rows[i].cells
            r0 = row[0].paragraphs[0].add_run(str(k))
            r0.bold = True
            row[1].paragraphs[0].add_run(str(v) if v is not None else "")

    # --- Tabla para Prescripciones ---
    def _add_prescripciones_table(self, doc, prescripciones: List[Dict[str, str]]):
        """
        Por cada prescripción crea una tabla de 2 columnas, excepto la fila 'Descripción'
        que ocupa el ancho completo (celdas fusionadas).
        """
        if not prescripciones:
            doc.add_paragraph("(sin prescripciones registradas)")
            return

        for m in prescripciones:
            # Crear tabla con encabezados
            table = doc.add_table(rows=1, cols=2)
            try:
                table.style = "Table Grid"
            except Exception:
                pass

            hdr = table.rows[0].cells
            run = hdr[0].paragraphs[0].add_run("Definición"); run.bold = True
            run = hdr[1].paragraphs[0].add_run("Descripción"); run.bold = True

            # Datos normalizados
            tipo = (m.get("tipo") or "").strip() or "Medida correctiva"
            desc = (m.get("descripcion") or "").strip() or "(sin descripción)"
            prioridad = (m.get("prioridad") or "").strip()
            plazo = (m.get("plazo") or "").strip()
            responsable = (m.get("responsable") or "").strip()

            # Filas normales a 2 columnas
            filas_normales = [
                ("Medida correctiva", tipo),
                ("Prioridad",         prioridad),
                ("Plazo",             plazo),
                ("Responsable",       responsable),
            ]
            for label, value in filas_normales:
                # Siempre incluir Plazo y Responsable, aunque value sea vacío.
                if label in ("Plazo", "Responsable") or value:
                    cells = table.add_row().cells
                    r0 = cells[0].paragraphs[0].add_run(label); r0.bold = True
                    cells[1].paragraphs[0].add_run(value or "")
            # Fila 'Descripción' en ancho completo (fusiona columnas)
            if desc:
                row = table.add_row()
                merged = row.cells[0].merge(row.cells[1])  # ← sin separación de columnas
                p = merged.paragraphs[0]
                r = p.add_run("Descripción: "); r.bold = True
                # agrega el texto (puede ser largo)
                p.add_run(desc)

            # Espacio entre tablas
            doc.add_paragraph()

    # ---------- API pública ----------
    def build(self, *, accidente, informe, resumen_texto: str = "") -> str:
        base = Path(settings.PROTECTED_MEDIA_ROOT) / "informes" / str(accidente.codigo_accidente)
        base.mkdir(parents=True, exist_ok=True)
        stem = f"{informe.codigo}_v{informe.version}"
        out_docx = base / f"{stem}.docx"

        empresa = self._leer_empresa_y_centro(accidente)
        trabajador = self._leer_trabajador(accidente)
        datos_acc = self._leer_datos_accidente(accidente)
        hechos = list(self._leer_hechos(accidente))
        prescripciones = list(self._leer_prescripciones(accidente))
        declaraciones = self._leer_declaraciones(accidente)
        documentos = self._leer_documentos(accidente)

        if not resumen_texto:
            relato_src = self._obtener_relato_para_resumen(accidente)
            resumen_texto = self._resumen_via_ia(relato_src) or \
                self._intentar_resumen_relato(accidente.accidente_id) or ""

        relato_final_text = self._obtener_relato_final(accidente)
        arbol_png = self._render_arbol_png(accidente)

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            h = doc.add_heading("Informe Técnico de Investigación", level=1)
            doc.add_paragraph()
            doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.add_run("RUT Empresa: ").bold = True
            p.add_run(str(empresa["rut_empresa"]))
            p = doc.add_paragraph()
            p.add_run("Fecha Accidente: ").bold = True
            p.add_run(self._fmt_date(getattr(accidente, "fecha_accidente", "")))
            p = doc.add_paragraph()
            p.add_run("Fecha Informe: ").bold = True
            p.add_run(self._fmt_date(getattr(informe, "fecha_informe", "")))

            doc.add_paragraph()
            doc.add_heading("1. Resumen del informe", level=2)
            # Separación visual entre el título y el texto siguiente
            doc.add_paragraph()
            # Resumen generado por IA
            resumen_ia = self._resumen_via_ia(accidente)
            doc.add_paragraph(resumen_ia or resumen_texto or "(sin resumen)")

            # ==== Sección 2 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass

            doc.add_paragraph()
            doc.add_heading("2. Metodología de análisis de causalidad", level=2)
            doc.add_paragraph()
            metodo = (
                "La metodología utilizada para la presente investigación se basa en la aplicación de un método denominado 'Árbol de Causas', el cual, entre otros, es promovido por la OIT y adoptado por el Ministerio de Salud de Chile para determinar la causalidad de los accidentes de origen laboral que deriven en consecuencias fatales o graves.\n"
                "Esta metodología permite, mediante un razonamiento lógico y secuencial, buscar de manera sistemática los hechos que han estado presentes en la ocurrencia del accidente y, como tal, facilitar la identificación de oportunidades de mejoramiento en los procesos de la empresa que se relacionan principalmente con la situación investigada.\n"
                "Elemento clave de esta metodología es descartar de forma categórica la incorporación de juicios de valor como elementos de análisis, considerándose solo aquellos elementos objetivos identificados y precisados durante el proceso investigativo —los hechos— que en definitiva corresponden a las causas que permitieron la ocurrencia del accidente. Estos hechos se representan gráficamente, lo cual facilita reconocer probables intervenciones y proponer acciones de mejora.\n"
                "Es necesario indicar que el modelo utilizado no permite, bajo ningún concepto o circunstancia, establecer o determinar culpables o responsables del accidente, sino solo facilitar la identificación de oportunidades de mejoramiento en los procesos de la empresa."
            )
            # Separar por saltos de línea y crear un párrafo por cada sección
            for parte in metodo.splitlines():
                if parte.strip():
                    doc.add_paragraph(parte.strip())
                    # agregar salto de línea adicional entre párrafos
                    doc.add_paragraph()
                else:
                    # mantener párrafo en blanco si hay líneas vacías
                    doc.add_paragraph()

            # ==== Sección 3 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass

            doc.add_paragraph()
            doc.add_heading("3. Antecedentes", level=2)
            doc.add_paragraph()

            doc.add_heading("3.1. Información de la Empresa y Centro de Trabajo", level=3)
            doc.add_paragraph()
            self._add_kv_table(doc, [
                ("Razón Social", empresa["razon_social"]),
                ("RUT Empresa", empresa["rut_empresa"]),
                ("Actividad Económica", empresa["actividad"]),
                ("Dirección Empresa", empresa["direccion_empresa"]),
                ("Teléfono Empresa", empresa["telefono_empresa"]),
                ("Representante Legal", empresa["representante_legal"]),
                ("Región", empresa["region"]),
                ("Comuna", empresa["comuna"]),
                ("Centro de Trabajo", empresa["centro_trabajo"]),
                ("Dirección Centro", empresa["direccion_centro"]),
            ])
            doc.add_paragraph()
            doc.add_heading("3.2. Datos del Trabajador", level=3)
            self._add_kv_table(doc, [
                ("Nombre Completo", trabajador["nombre"]),
                ("RUT Trabajador", trabajador["rut"]),
                ("Fecha de Nacimiento", trabajador["fecha_nacimiento"]),
                ("Edad", trabajador["edad"]),
                ("Nacionalidad", trabajador["nacionalidad"]),
                ("Estado Civil", trabajador["estado_civil"]),
                ("Tipo de Contrato", trabajador["tipo_contrato"]),
                ("Antigüedad en la empresa", trabajador["antiguedad_empresa"]),
                ("Cargo", trabajador["cargo"]),
                ("Antigüedad en el Cargo", trabajador["antiguedad_cargo"]),
                ("Domicilio", trabajador["domicilio"]),
            ])

            doc.add_paragraph()
            doc.add_heading("3.3. Datos del Accidente", level=3)
            self._add_kv_table(doc, [
                ("Fecha", datos_acc["fecha"]),
                ("Hora", datos_acc["hora"]),
                ("Lugar", datos_acc["lugar"]),
                ("Tipo", datos_acc["tipo"]),
                ("Naturaleza Lesión", datos_acc["naturaleza"]),
                ("Parte Afectada", datos_acc["parte_afectada"]),
                ("Tarea Ejecutada", datos_acc["tarea"]),
                ("Operación", datos_acc["operacion"]),
                ("Daños a Personas", datos_acc["danos_personas"]),
                ("Daños a Propiedad", datos_acc["danos_propiedad"]),
                ("Pérdidas en Proceso", datos_acc["perdidas_proceso"]),
            ])

            # ==== Sección 4 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass

            doc.add_paragraph()
            doc.add_heading("4. Principales fuentes de información", level=2)
            doc.add_paragraph()
            doc.add_paragraph("• Declaraciones y entrevistas")
            doc.add_paragraph("• Documentos adjuntos y fotografías")

            if declaraciones:
                doc.add_paragraph()
                doc.add_heading("4.1 Declaraciones (resumen)", level=3)
                doc.add_paragraph()
                for d in declaraciones:
                    doc.add_paragraph(f"- [{d['tipo']}] {d['nombre']} ({d['cargo']}) — {d['rut']}")
                    if d["texto"]:
                        doc.add_paragraph(f"  “{d['texto']}”")

            if documentos:
                doc.add_paragraph()
                doc.add_heading("4.2 Documentos/Fotografías", level=3)
                doc.add_paragraph()
                for d in documentos:
                    doc.add_paragraph(f"- {d['nombre']} ({d['mime']}) — {d['fecha']}")


            # ==== Sección 5 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass
            doc.add_paragraph()
            doc.add_heading("5. Descripción del Accidente", level=2)
            doc.add_paragraph()
            desc = (relato_final_text or self._first_non_empty(
                self._getattr_safe(accidente, "contexto"),
                self._getattr_safe(accidente, "circunstancias"),
                "",
            ))
            if desc:
                # separar por saltos de línea y crear un párrafo por cada línea (mantener líneas vacías)
                for linea in str(desc).splitlines():
                    p = doc.add_paragraph(linea if linea.strip() else "")
                    try:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        pf = p.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                    except Exception:
                        pass
            else:
                p = doc.add_paragraph("(sin descripción)")
                try:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    pass


            # ==== Sección 6 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass
            doc.add_paragraph()
            doc.add_heading("6. Principales hechos identificados", level=2)
            doc.add_paragraph()
            if hechos:
                for h in hechos:
                    doc.add_paragraph(f"{h}")
            else:
                doc.add_paragraph("(sin hechos registrados)")


            # ==== Sección 8 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass
            doc.add_paragraph()
            doc.add_heading("7. Árbol de Causas", level=2)
            doc.add_paragraph()
            if arbol_png and Path(arbol_png).exists():
                doc.add_picture(arbol_png, width=Inches(6.0))
                pic_par = doc.paragraphs[-1]
                pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                doc.add_paragraph("Figura 1. Árbol de Causas generado a partir de los hechos registrados.")
            else:
                doc.add_paragraph(
                    "No fue posible insertar la imagen del Árbol de Causas. "
                    "Se adjunta el archivo DOT en la carpeta del informe (si estuvo disponible)."
                )


            # ==== Sección 2 en hoja nueva
            try:
                doc.add_page_break()
            except Exception:
                pass
            # ======= Prescripciones en TABLA =======
            doc.add_paragraph()
            doc.add_heading("8. Prescripciones", level=2)
            doc.add_paragraph()
            if prescripciones:
                self._add_prescripciones_table(doc, prescripciones)
            else:
                doc.add_paragraph("(sin prescripciones registradas)")
            # =======================================
            
            # ===== Investigador responsable (desde accidente.usuario_asignado_id) =====
            try:
                # 1) Resolver el ID del usuario asignado (preferir FK en accidente)
                usuario_id = None
                try:
                    # si la FK viene materializada
                    usuario_id = getattr(getattr(accidente, "usuario_asignado", None), "id", None)
                except Exception:
                    usuario_id = None
                if not usuario_id:
                    usuario_id = getattr(accidente, "usuario_asignado_id", None)

                investigador = None
                if usuario_id:
                    # 2) Cargar modelo de usuario: primero users.Users, si no, AUTH_USER_MODEL
                    UserModel = (
                        self._get_model_if_exists("users", "Users") or
                        (lambda app_model: self._get_model_if_exists(*app_model.split(".")))(settings.AUTH_USER_MODEL)
                    )
                    if UserModel:
                        investigador = UserModel.objects.filter(pk=usuario_id).first()

                def _pick(obj, *names):
                    """Obtiene el primer atributo no vacío; soporta rutas con puntos."""
                    for n in names:
                        v = self._getattr_safe(obj, n, "")
                        if v:
                            return str(v)
                    return ""

                # 3) Extraer nombre(s) y cargo desde posibles campos
                nombre  = _pick(investigador, "nombre", "first_name")
                apepat  = _pick(investigador, "apepat", "last_name")
                apemat  = _pick(investigador, "apemat")
                cargo   = _pick(investigador, "Cargo", "cargo", "cargo.nombre", "position", "title")

                partes = [p for p in (nombre, apepat, apemat) if p]
                nombre_completo = " ".join(partes).strip()

                if nombre_completo or cargo:
                    doc.add_paragraph()
                    doc.add_paragraph()
                    # Pie de firma: nombre, cargo y etiqueta "Investigador responsable", todo centrado
                    doc.add_paragraph()
                    p_name = doc.add_paragraph()
                    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_name = p_name.add_run(nombre_completo or "")
                    r_name.bold = True

                    p_cargo = doc.add_paragraph()
                    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cargo.add_run(cargo or "")

                    # Línea indicadora debajo, centrada
                    p_label = doc.add_paragraph()
                    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_label.add_run("Investigador responsable")

            except Exception as e:
                log.warning("No fue posible insertar datos de investigador: %s", e)

            # ===== Estilo base del documento =====
            try:
                style = doc.styles["Normal"]
                style.font.name = "Calibri"
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                style.font.size = Pt(11)

                # Formato de párrafo por defecto: justificar y sin espaciado
                pfmt = style.paragraph_format
                pfmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pfmt.space_before = Pt(0)
                pfmt.space_after = Pt(0)

                # Asegurar que los párrafos dentro de tablas permanezcan alineados a la izquierda
                # (no queremos justificar el texto de tablas)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                pf = p.paragraph_format
                                pf.space_before = Pt(0)
                                pf.space_after = Pt(0)
            except Exception:
                pass
            doc.save(out_docx.as_posix())
            return out_docx.as_posix()

        except Exception as e:
            log.warning("Error generando DOCX (%s). Creo TXT fallback.", e)
            txt_path = (out_docx.with_suffix(".txt"))
            contenido = []
            A = contenido.append
            A("Informe Técnico de Investigación\n")
            A(f"RUT Empresa: {empresa['rut_empresa']}\n")
            A(f"Fecha Accidente: {self._fmt_date(getattr(accidente, 'fecha_accidente', ''))}\n")
            A(f"Fecha Informe: {self._fmt_date(getattr(informe, 'fecha_informe', ''))}\n\n")
            A("1. Resumen del informe\n")
            A((resumen_texto or "(sin resumen)") + "\n\n")
            A("2. Metodología de análisis de causalidad\nÁrbol de Causas (OIT/MINSAL)\n\n")
            A("3. Antecedentes\n")
            A(f"- Empresa: {empresa['razon_social']} ({empresa['rut_empresa']})\n")
            A(f"- Centro: {empresa['centro_trabajo']}\n")
            A("4. Hechos\n")
            for i, h in enumerate(hechos, 1):
                A(f"  {i}. {h}\n")
            A("5. Descripción del Accidente\n")
            desc_txt = (relato_final_text or
                        self._first_non_empty(self._getattr_safe(accidente, "contexto"),
                                              self._getattr_safe(accidente, "circunstancias"), ""))
            A((desc_txt or "(sin descripción)") + "\n\n")
            A("7. Árbol de Causas: ver archivo .dot si se generó.\n")
            # (Opcional: podrías listar prescripciones aquí en texto simple)
            txt_path.write_text("".join(contenido), encoding="utf-8")
            return txt_path.as_posix()
